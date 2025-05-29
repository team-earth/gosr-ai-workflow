"""
Script: json2doc.py

Purpose:
    This script generates a Microsoft Word (DOCX) outline from a hierarchical JSON tree (such as s.json or r.json)
    and a flat resource list (resources.json). It supports creating hyperlinks and bookmarks for resources,
    and can output resource details in a structured format for reporting or visualization.

Workflow:
    1. Loads configuration and resource list from the specified project directory.
    2. Loads the hierarchical tree structure (s.json or r.json).
    3. Builds a DOCX outline of the tree, with headings and resource details.
    4. Adds bookmarks and hyperlinks for resources.
    5. Saves the resulting DOCX file to the project directory.

Usage:
    python json2doc.py <project_subdirectory> --stage <stage_name>
    - <project_subdirectory> should contain config.yaml, s.json/r.json, and resources.json.
    - <stage_name> should be 'r' for resources or 's' for solutions.

Outputs:
    - <word_doc_title>.docx: Word document with the tree outline and resource details.

Dependencies:
    - Python 3.x
    - python-docx
    - pyyaml

See the project README for more details.
"""

import argparse
from docx import Document
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
import json
import os
import sys
import yaml
from html import escape

config = None
path = None

doc = Document()
global_resources_list = []
global_resources_dict = {}
hyperlink_list = []

def open_tree(filename):
    """
    Load a tree structure from a JSON file in the project directory.
    """
    if not isinstance(path, str) or not path:
        raise ValueError("The 'path' variable must be a non-empty string before calling open_tree.")
    with open(os.path.join(path, filename), "r", encoding='utf-8') as f:
        j = json.load(f)
    return j

def write_node(n, indent=0):
    """
    Recursively write a node and its children as a FreeMind XML outline (not used in DOCX output).
    """
    s = ""
    if type(n) == dict:
        for k, v in n.items():
            s += f'{"  "*indent}<node TEXT="{escape(k)}">\n'
            for c in v["children"]:
                s += write_node(c, indent + 2)
            s += f'{"  "*indent}</node>\n'
    elif type(n) == str:
        s += f'{"  "*indent}<node TEXT="{escape(n)}"></node>\n'
    return s

def write_section(doc, level, header, body):
    """
    Add a heading and paragraph to the DOCX document.
    """
    doc.add_heading(header.strip(), level)
    doc.add_paragraph(body.strip())

def write_resource(doc, level, d):
    """
    Add a resource as a bullet point with a hyperlink/bookmark in the DOCX document.
    """
    global global_resources_dict, hyperlink_list

    id = d["id"]
    if "dup" in global_resources_dict[id]:
        id = global_resources_dict[id]["dup"]

    hyperlink_para = doc.add_paragraph(style="List Bullet")

    r = global_resources_dict[id]
    if r["program"] != r["organization"]:
        run_text = f'{r["program"]} ({r["organization"]})'
    elif r["address"] != "N/A":
        run_text = f'{r["program"]} ({r["address"]})'
    else:
        run_text = f'{r["program"]}'
    bookmark_name = f"resource:{id}"

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    run = hyperlink_para.add_run(run_text)
    hyperlink.append(run._r)
    hyperlink_para._p.append(hyperlink)
    hyperlink_list.append(id)

def write_children(level, d):
    """
    Recursively write children nodes to the DOCX document, handling both headings and resources.
    """
    for c in d:
        if type(c) is dict:
            key = list(c.keys())[0]
            data = c[key]["data"]
            if key == "resource":
                write_resource(doc, level, data)
            else:
                if type(data) is dict:
                    header = data["title"]
                    body = data["description"]
                elif ":" in data:
                    (header, body) = data.split(":")
                else:
                    header = data
                    body = ""
                write_section(doc, level, header, body)
            if "children" in c[key]:
                write_children(level + 1, c[key]["children"])
        elif type(c) is str:
            data = c
            if ":" in data:
                (header, body) = data.split(":")
                write_section(doc, level, header, body)
            else:
                doc.add_paragraph(data, style="List Bullet")

def create_docx(j):
    """
    Create a DOCX outline from the tree structure.
    """
    k = list(j.keys())[0]
    doc.add_heading(j[k]["data"], 0)
    level = 1
    write_children(level, j[k]["children"])

def get_name_value(d):
    """
    Get the program/resource name from a resource dict.
    """
    possible_keys = ["name", "Name", "program_name", "program"]
    for key in possible_keys:
        if key in d:
            return d[key]
    print(f"No name key found in {d}")

def add_resource_paragraphs(resource):
    """
    Add a detailed paragraph for a resource, with a bookmark, to the DOCX document.
    """
    global hyperlink_list

    id = resource["id"]
    bookmark_name = f"resource:{id}"

    para = doc.add_paragraph()

    # Create a bookmark start element
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(id))
    bookmark_start.set(qn("w:name"), bookmark_name)
    para._p.append(bookmark_start)

    headers = ["program", "description", "organization", "address", "email", "website"]

    run = para.add_run(resource["program"])
    run.bold = True
    para.add_run("\n")
    para.add_run(f"{resource['description']} ")
    run = para.add_run("Organization: ")
    run.bold = True
    para.add_run(f'{resource["organization"]}; ')
    run = para.add_run("Address: ")
    run.bold = True
    para.add_run(f'{resource["address"]}; ')
    run = para.add_run("Email: ")
    run.bold = True
    para.add_run(f'{resource["email"]}; ')
    run = para.add_run("Website: ")
    run.bold = True
    run = para.add_run(resource["website"])

    # Create a bookmark end element
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(id))
    para._p.append(bookmark_end)

def add_resource(resource):
    """
    Add a resource's details as a table with a bookmark to the DOCX document.
    """
    global hyperlink_list

    name = resource["program"]
    id = resource["id"]
    bookmark_name = f"resource:{id}"

    para = doc.add_paragraph()
    table = doc.add_table(rows=1, cols=2)

    # Create a bookmark start element
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(id))
    bookmark_start.set(qn("w:name"), bookmark_name)
    para._p.addnext(bookmark_start)

    for k, v in resource.items():
        row_cells = table.add_row().cells
        row_cells[0].text = k.capitalize()
        row_cells[1].text = str(v)

    # Create a bookmark end element
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(id))
    para._p.addnext(bookmark_end)

def create_resources():
    """
    Add all resources from the global resource list to the DOCX document.
    """
    global global_resources_list
    global hyperlink_list

    doc.add_heading("Resources", 0)
    for id in hyperlink_list:
        add_resource_paragraphs(global_resources_dict[id])

def main():
    """
    Main entry point for the script.
    Loads config and resource list, builds the DOCX outline, and saves the result.
    """
    global config
    global path
    global global_resources_list
    global hyperlink_list

    parser = argparse.ArgumentParser(description="Generate a DOCX outline from a JSON tree and resource list.")
    parser.add_argument("path", help="Project directory containing config.yaml, s.json/r.json, and resources.json")
    parser.add_argument("--stage", choices=["r", "s"], required=True, help="Stage name: 'r' for resources, 's' for solutions (required)")
    args = parser.parse_args()

    path = args.path
    stage_name = args.stage

    with open(os.path.join(path, "config.yaml"), "r", encoding='utf-8') as file:
        config = yaml.safe_load(file)

    docx_title = config.get('word_doc_title', config['root_node_name'])
    docx_file_path = os.path.join(path, f"{docx_title}.docx")

    try:
        with open(docx_file_path, "a", encoding='utf-8'):
            pass
    except (IOError, OSError):
        print(f"File locked? Could not open {docx_file_path}")
        sys.exit(0)

    # If working with resources, load the resource list and build a lookup dict
    if stage_name == "r":
        with open(os.path.join(path, "resources.json"), "r", encoding='utf-8') as file:
            global_resources_list = json.load(file)
        for r in global_resources_list:
            global_resources_dict[r["id"]] = r

    # Load the tree and create the DOCX outline
    j = open_tree(f"{stage_name}.json")
    create_docx(j)

    # If working with resources, add detailed resource paragraphs
    if stage_name == "r":
        create_resources()

    doc.save(docx_file_path)
    return 0

if __name__ == "__main__":
    sys.exit(main())
