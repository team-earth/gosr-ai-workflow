import docx
import openai
from treelib import Node, Tree
import json
# import codecs
import re
import sys
import time
from html import escape
import os
import yaml
import csv

config = None
path = None

stage_name = 'r'
doc = docx.Document()

def open_tree(filename = f"{stage_name}.json"):
  with open(os.path.join(path,filename), 'r', encoding='utf-8') as f:
    j = json.load(f)
#   print(json.dumps(j, indent=2))
  return j

def write_node(n, indent=0):
    s = ""

    if type(n) == dict:
        for k, v in n.items():
            s += f'{"  "*indent}<node TEXT="{escape(k)}">\n'
            for c in v['children']:
                s += write_node(c, indent+2)
            s += f'{"  "*indent}</node>\n'
    elif type(n) == str:
        s += f'{"  "*indent}<node TEXT="{escape(n)}"></node>\n'

    return s
   
def write_section(doc, level, header, body):
    doc.add_heading(header.strip(), level)
    doc.add_paragraph(body.strip())

def write_children(level, d):
    for c in d:
        if type(c) is dict:
            key = list(c.keys())[0]
            text = c[key]["data"]
            # if 
            if ':' in text:
                (header,body) = text.split(':')
            else:
                header = text
                body = ""
            if key == 'resource':
                write_section(doc, level, "RESOURCES", text)
            else:
                write_section(doc, level, header, body)
            if "children" in c[key]:
                write_children(level+1, c[key]['children'])
        elif type(c) is str:
            text = c
            if ':' in text:
                (header,body) = text.split(':')
                write_section(doc, level, header, body)
            else:
                # write this solution as a list item
                doc.add_paragraph(text, style='List Bullet')

def create_docx(j):
    k = list(j.keys())[0]
    doc.add_heading(j[k]["data"],0)
    level = 1

    write_children(level, j[k]['children'])

def get_name_value(d):
    possible_keys = ['name', 'Name', 'program_name', 'program']
    for key in possible_keys:
        if key in d:
            return d[key]
    print(f"No name key found in {d}")

def add_resource(resource):
    name = get_name_value(resource)
    doc.add_heading(name, 1)
    keys = [
        "description", "organization", "address", "email", "website"
    ]
    for k in keys:
        doc.add_paragraph(f"{k}: {resource[k]}")


def create_resources(filename ):
  with open(os.path.join(path,filename), 'r', encoding='utf-8') as f:
    resources = json.load(f)
    doc.add_heading("Resources", 0)
  for resource in resources:
      add_resource(resource)



aliases = {
    "name": [
        "name",
        "Name",
        "program_name",
        "program",
        "event_name",
        "title",
        "Effort Name",
        "EffortName",
        "Effort",
        "effort_name",
        "project_name",
        "description",
        "effort_description",
        "EffortDescription",
        "Effort_Name",
        "ProgramName",
        "ProjectName",
        "effort"
    ],
    "description": ["description", "Description", "effort_description", "effort", "EffortDescription"],
    "organization": ["organization", "address", "Organization", "ImplementingOrganization", "OrganizationName"],
    "address": ["address", "Address"],
    "email": ["email", "Email", "email unavailable"],
    "website": ["website", "Website"]
}

def is_list_type(d):
    list_keys = ['effort_1', 'LondonEfforts']
    for list_key in list_keys:
        if list_key in d:
            return True
    return False

def get_value(key, d):
    for k in aliases[key]:
        if k in d:
            return d[k]

    if 'organizations' in d:
        if key == 'organization':
            return d['organizations'][0]['name']
        if key == 'address':
            return d['organizations'][0]['address']
        if key == 'email':
            return d['organizations'][0]['email']
        if key == 'website':
            return d['organizations'][0]['website']

    if key == 'name':
        return d['organization']['name']
    if key == 'address':
        return d['organization']['address']
    if key == 'email':
        return d['organization']['email']
    if key == 'website':
        return d['organization']['website']
    print(f"No key '{key}' found in {d}")
    sys.exit(1)

def normalize_resource_list():
    global resource_list
    keys = ["name", "description", "organization", "address", "email", "website"]
    for r in resource_list:
        elem = {}
        for k in keys:
            elem[k] = get_value(k, r)
        r = elem

rcount = 0
resource_keys = {}

def find_by_id(id):
    global resource_list
    for k in resource_list:
        if k["id"] == id:
            if "dup" in k:
                return find_by_id(k["dup"])
            else:
                return k
    return None


def get_all_resources(j, parent):
    global rcount, resource_list

    rs = []

    key = list(j.keys())[0]
    if key == "solution":
        if "children" in j[key]:
            for r in j[key]["children"]:
                # print(r)
                resource = find_by_id(r["resource"]["data"]["id"])
                solving_dict = parent["obstacle"]["data"]
                if type(solving_dict) == dict:
                    solving = solving_dict["title"].rstrip('. ')+'. '+solving_dict["description"].rstrip('. ')+'.'
                else:
                    solving = solving_dict
                
                solution_dict = j[key]["data"]
                if type(solution_dict) == dict:
                    solution = solution_dict["title"].rstrip('. ')+'. '+solution_dict["description"].rstrip('. ')+'.'
                else:
                    solution = solution_dict
                
                resource["solving"] = solving
                resource["solution"] = solution

                # fix URLs
                if "url_valid" in resource:
                    if resource["url_valid"] == False:
                        del resource["url_valid"]
                        resource["website"] = "N/A"
                    elif resource["url_valid"] == True:
                        del resource["url_valid"]
                    else:
                        resource["website"] = resource["url_valid"]
                        del resource["url_valid"]

                rs.append(resource)
        return rs
    else:
        for c in j[key]["children"]:
            rs.extend(get_all_resources(c,j))
        return rs
    
resource_list = []

def main():
    global config
    global path
    global resource_list

    if len(sys.argv) != 2:
        print(f'Usage: {sys.argv[0]} path')
        return 1

    path = sys.argv[1]
    with open(os.path.join(path,'config.yaml'), 'r', encoding='utf-8') as file:
        config = yaml.safe_load(file)

    with open(os.path.join(path, "resources.json"), "r", encoding='utf-8') as f:
        resource_list = json.load(f)

    # normalize_resource_list()

    with open(os.path.join(path,'r.json'), 'r', encoding='utf-8') as f:
        j = json.load(f)
 
    r_all = []
    unique = []

    for top_level_theme in j["goal"]["children"]:
        if "label" in top_level_theme["obstacle"]:
            theme = top_level_theme["obstacle"]["label"]
        else:
            theme = top_level_theme["obstacle"]["data"]["title"]
        r = get_all_resources(top_level_theme, None)

        # filename = theme.strip()
        # filename = re.sub(r'/', ', ', filename)
        # filename = re.sub(r':', ' - ', filename)
        # filename = re.sub(r'[^a-zA-Z0-9_.-]', ' ', filename)
        # filename = re.sub(r'  +', ' ', filename)
        filename = "out"

        data = r[0]
        # data['theme'] = theme

        os.makedirs(os.path.join(path,"CSV"), exist_ok=True)
        with open(os.path.join(path,"CSV",f"{filename}.csv"), "a", newline='\n', encoding='utf-8') as csvfile:
            writer = csv.DictWriter(csvfile, fieldnames=data.keys())
            writer.writeheader()
            writer.writerows(r)

        for d in r:
            if d["id"] not in unique:
                unique.append(d["id"])
                d_copy = d.copy()
                d_copy.update({'category': theme})
                r_all.append(d_copy)

    with open(os.path.join(path,"mailing_list.csv"), "w", newline='\n', encoding='utf-8') as csvfile:
        header = list(r_all[0].keys())
        writer = csv.DictWriter(csvfile, fieldnames=header)
        writer.writeheader()
        writer.writerows(r_all)
    return 0

if __name__ == '__main__':
    sys.exit(main())