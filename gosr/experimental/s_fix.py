from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

import openai
from treelib import Node, Tree
import json

# import codecs
# import re
import sys
import time
from html import escape
import os
import yaml

config = None
path = None

stage_name = "r"
doc = Document()
global_resources_list = []
global_resources_dict = {}
hyperlink_list = []

new_j = {}

def open_tree(filename):
    with open(os.path.join(path, filename), "r", encoding='utf-8') as f:
        j = json.load(f)
    #   print(json.dumps(j, indent=2))
    return j


def write_node(n, indent=0):
    s = ""

    if type(n) == dict:
        for k, v in n.items():
            s += f'{"  "*indent}<node TEXT="{escape(k)}">\n'
            for c in v["children"]:
                s += write_node(c, indent + 2)
            s += f'{"  "*indent}</node>\n'
    elif type(n) == str:
        s += f'{"  "*indent}<node TEXT="{escape(n)}"></node>\n'

    return s


def write_section(doc, level, header, body):
    doc.add_heading(header.strip(), level)
    doc.add_paragraph(body.strip())


def write_resource(doc, level, d):
    global global_resources_dict, hyperlink_list

    id = d["id"]
    if "dup" in global_resources_dict[id]:
        id = global_resources_dict[id]["dup"]

    hyperlink_para = doc.add_paragraph(style='List Bullet')

    r = global_resources_dict[id]
    if r["program"] != r["organization"]:
        run_text = f'{r["program"]} ({r["organization"]})'
    elif r["address"] != 'N/A':
        run_text = f'{r["program"]} ({r["address"]})'
    else:
        run_text = f'{r["program"]}'
    bookmark_name = f"resource:{id}"

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    run = hyperlink_para.add_run(run_text)
    hyperlink.append(run._r)
    hyperlink_para._p.append(hyperlink)
    hyperlink_list.append(id)


def write_children(level, d):
    count = 0
    for c in d:
        if type(c) is dict:
            key = list(c.keys())[0]
            data = c[key]["data"]
            # if
            if ":" in data:
                (header, body) = data.split(":")
            else:
                header = data
                body = ""
            if key == "resource":
                write_resource(doc, level, data)
            else:
                write_section(doc, level, header, body)
            if "children" in c[key]:
                write_children(level + 1, c[key]["children"])
        elif type(c) is str:
            data = c
            if ":" in data:
                (header, body) = data.split(":")
                write_section(doc, level, header, body)
            else:
                # write this solution as a list item
                doc.add_paragraph(data, style="List Bullet")
        
        # return prematurely
        # count += 1
        # if count > 2:
        #     return

# def check_children(n_list):
#     children = n_list[-1]
#     assert(type(children) == list)
#     for c in children:
#         assert(type(c) == dict)
#         assert(len(c.keys()) == 1)
#         k = next(iter(c.keys()))
#         check_node(c[k])
#         assert(k in ['obstacle', 'solution', 'resource'])
#         if "children" in c[k]:
#             check_children(n_list + )
        # if k == 'obstacle':
        #     analyze_obstacle()
        # elif k == 'solution':
        #     analyze_solution()
        # elif k == 'resource':
        #     analyze_resource()
        # else:
        #     sys.exit(1)

count = 0
def fix_double_obstacle(n):
    global count
    k = next(iter(n.keys()))
    if k == 'obstacle':
        if len(n[k]["children"]) == 1:
            cn = n[k]["children"][0]
            ck = next(iter(cn.keys()))
            if cn[ck]["data"] == n[k]["data"]:
                for x in [(k,n,"parent"), (ck,cn,"child")]:
                    print(x[2], "\t", x[0], len(x[1][x[0]]["children"]), "children", x[1][x[0]]["data"] )
                count += 1
                print(f"Fixing {count:04g}")
                n[k]["children"] = cn[k]["children"]
    if "children" in n[k]:
        assert(type(n[k]["children"]) == list)
        for c in n[k]["children"]:
            fix_double_obstacle(c)

rcount = 0    
def resource_ids(n):
    global rcount
    k = next(iter(n.keys()))
    if k == 'resource':
        assert("children" not in n[k])
        assert("data" in n[k])
        if type(n[k]["data"]) == str:
            if ("id" in n[k]):
                d = {
                    "program": n[k]["data"],
                    "id": n[k]["id"]
                }
                del n[k]["id"]
                del n[k]["data"]
                n[k]["data"] = d
                rcount += 1
                print(f"Fixed {rcount:04g}")
        if type(n[k]["data"]) == dict:
            assert("program" in n[k]["data"])
            assert("id" in n[k]["data"])
    if "children" in n[k]:
        assert(type(n[k]["children"]) == list)
        for c in n[k]["children"]:
            resource_ids(c)

def swap(p):
    # p = p_list[-1]
    assert(type(p) == dict)

    k = next(iter(p.keys()))
    assert(len(p.keys()) == 1)
    assert(k in ['goal', 'obstacle', 'solution', 'resource'])
    n = p[k]
    if "data" not in n:
        print("Key 'data' not among keys", n.keys())
    data = n["data"]
    if type(data) is dict:
        if "title" in data and "description" in data:
            title = data["title"]
            description = data["description"]
            if len(title) > len(description):
                print("Case", title, description)
                data["description"] = title
                data["title"] = description
        
    # assert(type(n["data"]) == str)
    # if len(n.keys()) == 3:
    #     assert("label" in n)
    #     assert(type(n["label"]) == str)

    # if len(p_list) > 1:
    #     parent = p_list[-2]
    #     parent_key = next(iter(parent.keys()))
    #     if k == 'solution':
    #         if parent_key == 'solution':
    #                 print("Solution parent and child", parent, p)
    #     if n["data"] == parent[parent_key]["data"]:
    #         print("Parent and child duplicate data:")
    #         for x in [(k,n,"child"), (parent_key,parent[parent_key], "parent")]:
    #             print(x[2], "\t", x[0], len(x[1]["children"]), "children", x[1]["data"] )
    #         print()
    if "children" in n:
        assert(type(n["children"]) == list)
        for c in n["children"]:
            # c_list = list(p_list)
            # c_list.append(c)
            swap(c)

def analyze_solutions(j):
    fix_double_obstacle(j)
    print("Rerunning analyze_solutions")
    fix_double_obstacle(j)

def analyze_resources(j):
    resource_ids(j)
    print("Rerunning analyze_resources")
    resource_ids(j)

def main():
    global config
    global path
    global global_resources_list
    global hyperlink_list

    if len(sys.argv) != 2:
        print(f"Usage: {sys.argv[0]} path")
        return 1

    path = sys.argv[1]
    with open(os.path.join(path, "config.yaml"), "r", encoding='utf-8') as file:
        config = yaml.safe_load(file)

    # docx_file_path = os.path.join(path, f"{config['root_node_name']}.docx")

    # try:
    #     with open(docx_file_path, "a", encoding='utf-8'):
    #         pass
    # except (IOError, OSError):
    #     print(f"File locked? Could not open {docx_file_path}")
    #     sys.exit(0)

    # with open(os.path.join(path, "resources.json"), "r", encoding='utf-8') as file:
    #     global_resources_list = json.load(file)
    # for r in global_resources_list:
    #     global_resources_dict[r["id"]] = r

    j = open_tree("s-orig.json")

    print("Run 1")
    swap(j)
    print("Run 2")
    swap(j)


    with open(os.path.join(path, "s.json"), "w", encoding='utf-8') as f:
        json.dump(j, f)

    # doc.save(docx_file_path)
    return 0


if __name__ == "__main__":
    sys.exit(main())
