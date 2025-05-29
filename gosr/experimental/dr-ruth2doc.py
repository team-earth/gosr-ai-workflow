from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

import openai
from treelib import Node, Tree
import json

# import codecs
# import re
import sys
import time
from html import escape
import os
import yaml

config = None
path = None

stage_name = "s"
doc = Document()
global_resources_list = []
global_resources_dict = {}
hyperlink_list = []


def open_tree(filename=f"{stage_name}.json"):
    with open(os.path.join(path, filename), "r", encoding='utf-8') as f:
        j = json.load(f)
        # print(json.dumps(j, indent=2))
    return j


def write_node(n, indent=0):
    s = ""

    if type(n) == dict:
        for k, v in n.items():
            s += f'{"  "*indent}<node TEXT="{escape(k)}">\n'
            for c in v["children"]:
                s += write_node(c, indent + 2)
            s += f'{"  "*indent}</node>\n'
    elif type(n) == str:
        s += f'{"  "*indent}<node TEXT="{escape(n)}"></node>\n'

    return s


def write_section(doc, level, header, body):
    doc.add_heading(header.strip(), level)
    doc.add_paragraph(body.strip())


def write_resource(doc, level, d):
    global global_resources_dict, hyperlink_list

    id = d["id"]
    if "dup" in global_resources_dict[id]:
        id = global_resources_dict[id]["dup"]

    hyperlink_para = doc.add_paragraph(style="List Bullet")

    r = global_resources_dict[id]
    if r["program"] != r["organization"]:
        run_text = f'{r["program"]} ({r["organization"]})'
    elif r["address"] != "N/A":
        run_text = f'{r["program"]} ({r["address"]})'
    else:
        run_text = f'{r["program"]}'
    bookmark_name = f"resource:{id}"

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    run = hyperlink_para.add_run(run_text)
    hyperlink.append(run._r)
    hyperlink_para._p.append(hyperlink)
    hyperlink_list.append(id)


def write_children(level, d):
    count = 0
    for c in d:
        if type(c) is dict:
            key = list(c.keys())[0]
            data = c[key]["data"]
            if key == "resource":
                write_resource(doc, level, data)
            else:
                if type(data) is dict:
                    header = data["title"]
                    body = data["description"]
                elif ":" in data:
                    (header, body) = data.split(":")
                else:
                    header = data
                    body = ""
                write_section(doc, level, header, body)
            if "children" in c[key]:
                write_children(level + 1, c[key]["children"])
        elif type(c) is str:
            data = c
            if ":" in data:
                (header, body) = data.split(":")
                write_section(doc, level, header, body)
            else:
                # write this solution as a list item
                doc.add_paragraph(data, style="List Bullet")

        # return prematurely
        # count += 1
        # if count > 2:
        #     return
def write_program(p_id):
    program = global_resources_dict[str(p_id)]
    # doc.add_heading(program["program_name"].strip(), 3)
    # doc.add_paragraph(program["program_description"].strip())
    # doc.add_paragraph(', '.join([
    #     program["organization_name"],
    #     program["organization_description"],
    #     program["address"],
    #     program["phone"],
    #     program["website"]        
    # ]))
    
    para = doc.add_paragraph()

    # Create a bookmark start element
    # bookmark_start = OxmlElement("w:bookmarkStart")
    # bookmark_start.set(qn("w:id"), str(id))
    # bookmark_start.set(qn("w:name"), bookmark_name)

    # para._p.append(bookmark_start) –

    run = para.add_run(f'{program["program_name"].strip()} – ')
    run.bold = True
    run = para.add_run(f'{program["program_description"].strip()} ')
    run = para.add_run("Organization: ")
    run.bold = True
    para.add_run(f'{program["organization_name"]}. ')
    para.add_run(f'{program["organization_description"]} ')
    run = para.add_run("Address: ")
    run.bold = True
    para.add_run(f'{program["address"]}; ')
    if program["phone"].strip() != "":
        run = para.add_run("Phone: ")
        run.bold = True
        para.add_run(f'{program["phone"]}; ')
    run = para.add_run("Website: ")
    run.bold = True
    run = para.add_run(program["website"])



def write_section(section):
    doc.add_heading(section["section"].strip(), 2)
    doc.add_paragraph(section["description"].strip())
    doc.add_paragraph(f'Keywords: {", ".join(section["topics"])}')
    for p_id in section["programs"]:
        write_program(p_id)

def write_chapter(chapter, sections):
    doc.add_heading(chapter.strip(), 1)
    # doc.add_paragraph(body.strip())
    for section in sections["sections"]:
        write_section(section)

def create_docx(chapters):
    for chapter, sections in chapters.items():
        write_chapter(chapter, sections)


def get_name_value(d):
    possible_keys = ["name", "Name", "program_name", "program"]
    for key in possible_keys:
        if key in d:
            return d[key]
    print(f"No name key found in {d}")


def add_resource_paragraphs(resource):
    global hyperlink_list

    id = resource["id"]
    bookmark_name = f"resource:{id}"

    para = doc.add_paragraph()

    # Create a bookmark start element
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(id))
    bookmark_start.set(qn("w:name"), bookmark_name)

    para._p.append(bookmark_start)

    headers = ["program", "description", "organization", "address", "email", "website"]

    run = para.add_run(resource["program"])
    run.bold = True
    para.add_run("\n")
    para.add_run(f"{resource['description']} ")
    run = para.add_run("Organization: ")
    run.bold = True
    para.add_run(f'{resource["organization"]}; ')
    run = para.add_run("Address: ")
    run.bold = True
    para.add_run(f'{resource["address"]}; ')
    run = para.add_run("Email: ")
    run.bold = True
    para.add_run(f'{resource["email"]}; ')
    run = para.add_run("Website: ")
    run.bold = True
    run = para.add_run(resource["website"])

    # para.add_run('\n')

    # run.bold = True

    # Create a bookmark end element
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(id))

    # Insert the bookmark start and end elements into the paragraph
    para._p.append(bookmark_end)


def add_resource(resource):
    global hyperlink_list

    name = resource["program"]
    id = resource["id"]
    bookmark_name = f"resource:{id}"

    para = doc.add_paragraph()
    table = doc.add_table(rows=1, cols=2)

    # Create a bookmark start element
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(id))
    bookmark_start.set(qn("w:name"), bookmark_name)

    para._p.addnext(bookmark_start)

    for k, v in resource.items():
        row_cells = table.add_row().cells
        row_cells[0].text = k.capitalize()
        row_cells[1].text = str(v)

    # Create a bookmark end element
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(id))

    # Insert the bookmark start and end elements into the paragraph
    para._p.addnext(bookmark_end)


def create_resources():
    global global_resources_list
    global hyperlink_list

    doc.add_heading("Resources", 0)
    for id in hyperlink_list:
        add_resource_paragraphs(global_resources_dict[id])


def main():
    global config
    global path
    global global_resources_dict
    global hyperlink_list

    if len(sys.argv) != 2:
        print(f"Usage: {sys.argv[0]} path")
        return 1

    path = sys.argv[1]
    with open(os.path.join(path, "config.yaml"), "r", encoding='utf-8') as file:
        config = yaml.safe_load(file)

    docx_title = config.get('word_doc_title', config['root_node_name'])
    docx_file_path = os.path.join(path, f"{docx_title}.docx")

    try:
        with open(docx_file_path, "a", encoding='utf-8'):
            pass
    except (IOError, OSError):
        print(f"File locked? Could not open {docx_file_path}")
        sys.exit(0)

    with open(os.path.join(path, "programs.json"), "r", encoding='utf-8') as file:
        global_resources_dict = json.load(file)

    with open(os.path.join(path, "solutions.json"), "r", encoding='utf-8') as file:
        solutions = json.load(file)

    chapters = {}
    for d in solutions:
        chapter = d["chapter"]
        if chapter not in chapters.keys():
            chapters[chapter] = { "sections": []}
        chapters[chapter]["sections"].append(
            {key: d[key] for key in d.keys() if key not in ["chapter"]}
        )

    create_docx(chapters)

    # create_resources()

    doc.save(docx_file_path)
    return 0


if __name__ == "__main__":
    sys.exit(main())
