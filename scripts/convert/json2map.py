"""
Script: json2map.py

Purpose:
    This script processes a hierarchical JSON tree (such as r.json) and a flat resource list (resources.json)
    to generate a new JSON file with fixed solution/resource mappings. It can also generate a DOCX outline
    of the tree and resources for reporting or visualization.

Workflow:
    1. Loads configuration and resource list from the specified project directory.
    2. Loads the hierarchical tree structure (r.json).
    3. Fixes solution/resource mappings in the tree (handles "double solutions").
    4. Writes the fixed tree to r-fixed-double-solutions.json.
    5. (Optional, commented out) Can generate DOCX outlines or resource summaries.

Usage:
    python json2map.py <project_subdirectory>
    - <project_subdirectory> should contain config.yaml, r.json, and resources.json.

Outputs:
    - r-fixed-double-solutions.json: Tree with corrected solution/resource mappings.
    - (Optional) DOCX file with outline or resource summaries.

Dependencies:
    - Python 3.x
    - python-docx
    - treelib
    - openai
    - pyyaml

See the project README for more details.
"""

import docx
import json
import sys
import os
import yaml
import html
import argparse

config = None
path = ""

def escape(s):
    """
    Escape special characters for XML/HTML output.
    """
    return html.escape(str(s), quote=True)

doc = docx.Document()

def open_tree(filename):
    """
    Load a tree structure from a JSON file in the project directory.
    """
    with open(os.path.join(path, filename), 'r', encoding='utf-8') as f:
        j = json.load(f)
    return j

def write_node(n, indent=0):
    """
    Recursively write a node and its children as a FreeMind XML outline.
    """
    s = ""
    if type(n) == dict:
        for k, v in n.items():
            s += f'{"  "*indent}<node TEXT="{escape(k)}">\n'
            for c in v['children']:
                s += write_node(c, indent+2)
            s += f'{"  "*indent}</node>\n'
    elif type(n) == str:
        s += f'{"  "*indent}<node TEXT="{escape(n)}"></node>\n'
    return s

def write_section(doc, level, header, body):
    """
    Add a heading and paragraph to the DOCX document.
    """
    doc.add_heading(header.strip(), level)
    doc.add_paragraph(body.strip())

def write_children(level, d):
    """
    Recursively write children nodes to the DOCX document.
    """
    for c in d:
        if type(c) is dict:
            key = list(c.keys())[0]
            text = c[key]["data"]
            if ':' in text:
                (header, body) = text.split(':')
            else:
                header = text
                body = ""
            if key == 'resource':
                write_section(doc, level, "RESOURCES", text)
            else:
                write_section(doc, level, header, body)
            if "children" in c[key]:
                write_children(level+1, c[key]['children'])
        elif type(c) is str:
            text = c
            if ':' in text:
                (header, body) = text.split(':')
                write_section(doc, level, header, body)
            else:
                doc.add_paragraph(text, style='List Bullet')

def create_docx(j):
    """
    Create a DOCX outline from the tree structure.
    """
    k = list(j.keys())[0]
    doc.add_heading(j[k]["data"], 0)
    level = 1
    write_children(level, j[k]['children'])

def get_name_value(d):
    """
    Get the program/resource name from a resource dict.
    """
    possible_keys = ['name', 'Name', 'program_name', 'program']
    for key in possible_keys:
        if key in d:
            return d[key]
    print(f"No name key found in {d}")

def add_resource(resource):
    """
    Add a resource's details to the DOCX document.
    """
    name = get_name_value(resource)
    if not isinstance(name, str) or name is None:
        name = "Unknown Resource"
    doc.add_heading(name, 1)
    keys = [
        "description", "organization", "address", "email", "website"
    ]
    for k in keys:
        doc.add_paragraph(f"{k}: {resource[k]}")

def create_resources(filename):
    """
    Add all resources from a JSON file to the DOCX document.
    """
    with open(os.path.join(path, filename), 'r', encoding='utf-8') as f:
        resources = json.load(f)
        doc.add_heading("Resources", 0)
    for resource in resources:
        add_resource(resource)

def is_list_type(d):
    """
    Heuristic to check if a dict contains a list of efforts/resources.
    """
    list_keys = ['effort_1', 'LondonEfforts']
    for list_key in list_keys:
        if list_key in d:
            return True
    return False

# Define aliases for standard keys used in resources
aliases = {
    'name': ['name', 'Name', 'program_name', 'program'],
    'organization': ['organization', 'Organization', 'org'],
    'address': ['address', 'Address', 'location'],
    'email': ['email', 'Email'],
    'website': ['website', 'Website', 'url', 'URL']
}

def get_value(key, d):
    """
    Get a value for a standard key from a resource dict, handling nested orgs.
    """
    for k in aliases[key]:
        if k in d:
            return d[k]
    if 'organizations' in d:
        if key == 'organization':
            return d['organizations'][0]['name']
        if key == 'address':
            return d['organizations'][0]['address']
        if key == 'email':
            return d['organizations'][0]['email']
        if key == 'website':
            return d['organizations'][0]['website']
    if key == 'name':
        return d['organization']['name']
    if key == 'address':
        return d['organization']['address']
    if key == 'email':
        return d['organization']['email']
    if key == 'website':
        return d['organization']['website']
    print(f"No key '{key}' found in {d}")
    sys.exit(1)

rcount = 0
resource_keys = {}

def find_in_resource_list(name):
    """
    Find the resource ID for a given program/resource name.
    """
    global resource_list
    for k in resource_list:
        if k["name"] == name:
            return k["id"]
    return None

def fix_double_solutions(j):
    """
    Fixes "double solution" nodes in the tree by splitting/merging resource nodes.
    Updates the tree in-place.
    """
    global rcount, resource_keys
    key = list(j.keys())[0]
    if key != "solution":
        if "children" in j[key]:
            for c in j[key]["children"]:
                fix_double_solutions(c)
    else:
        r_dict = j["solution"]["children"][0]["solution"]["children"][0]
        if r_dict["resource"]["data"] == "":
            del j["solution"]["children"]
        else:  # r_dict is good
            del j["solution"]["children"][0]
            data = r_dict["resource"]["data"]
            if '|' in data:
                resources = data.split('|')
            else:
                resources = data.split(", ")
            for r in resources:
                id = find_in_resource_list(r)
                j["solution"]["children"].append({
                    "resource": {
                        "id": id,
                        "data": r
                    }
                })

resource_list = []

def main():
    """
    Main entry point for the script.
    Loads config and resource list, fixes double solutions in the tree, and saves the result.
    """
    global config
    global path
    global resource_list

    parser = argparse.ArgumentParser(description="Fix solution/resource mappings in a tree JSON file.")
    parser.add_argument("path", help="Project directory containing config.yaml, r.json/s.json, and resources.json")
    parser.add_argument("--stage", choices=["r", "s"], required=True, help="Stage name: 'r' for resources, 's' for solutions (required)")
    args = parser.parse_args()

    path = args.path
    stage_name = args.stage

    with open(os.path.join(path, 'config.yaml'), 'r', encoding='utf-8') as file:
        config = yaml.safe_load(file)

    with open(os.path.join(path, "resources.json"), "r", encoding='utf-8') as f:
        resource_list = json.load(f)

    for r in resource_list:
        if "dup" in r:
            continue
        resource_keys[r["program"]] = r["id"]

    with open(os.path.join(path, f'{stage_name}.json'), 'r', encoding='utf-8') as f:
        j = json.load(f)

    fix_double_solutions(j)
    with open(os.path.join(path, f"{stage_name}-fixed-double-solutions.json"), "w", encoding='utf-8') as f:
        json.dump(j, f)

    return 0

if __name__ == '__main__':
    sys.exit(main())